import streamlit as st
import os
import sys
from docx import Document
from openai import OpenAI
import re
from typing import List, Tuple, Optional
from dotenv import load_dotenv
from config import Config
import fitz  # pymupdf

# Detect if running as standalone executable (prod mode)
# PyInstaller creates a 'frozen' attribute when bundled as exe
IS_PROD_MODE = getattr(sys, 'frozen', False)

# Load environment variables (only in dev mode)
if not IS_PROD_MODE:
    load_dotenv()

# Initialize session state for API key (prod mode)
if 'openai_api_key' not in st.session_state:
    st.session_state.openai_api_key = ""

# Supported languages list (alphabetical order)
SUPPORTED_LANGUAGES = [
    "Arabic",
    "Bengali",
    "Bulgarian",
    "Catalan",
    "Chinese (Simplified)",
    "Chinese (Traditional)",
    "Croatian",
    "Czech",
    "Danish",
    "Dutch",
    "English",
    "Estonian",
    "Finnish",
    "French",
    "German",
    "Greek",
    "Hebrew",
    "Hindi",
    "Hungarian",
    "Icelandic",
    "Indonesian",
    "Irish",
    "Italian",
    "Japanese",
    "Kazakh",
    "Korean",
    "Latvian",
    "Lithuanian",
    "Malay",
    "Maltese",
    "Norwegian",
    "Persian (Farsi)",
    "Polish",
    "Portuguese",
    "Romanian",
    "Russian",
    "Serbian",
    "Slovak",
    "Slovenian",
    "Spanish",
    "Swedish",
    "Tagalog",
    "Tamil",
    "Thai",
    "Turkish",
    "Ukrainian",
    "Urdu",
    "Vietnamese",
    "Welsh",
    "Yiddish"
]

# Page configuration
st.set_page_config(
    page_title="Global Translator: AI Translation for Large Text Files",
    page_icon="🌍",
    layout="wide"
)

# Initialize OpenAI client
def get_openai_client():
    """
    Get OpenAI client with API key from appropriate source:
    - Dev mode: from .env file
    - Prod mode: from session state (user input in UI)
    """
    if IS_PROD_MODE:
        # Production mode: get API key from session state
        api_key = st.session_state.get('openai_api_key', '')
        if not api_key or api_key.strip() == '':
            st.error("⚠️ Please enter your OpenAI API key in the Settings section (sidebar)")
            st.stop()
    else:
        # Development mode: get API key from .env file
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            st.error("Please set OPENAI_API_KEY in your .env file")
            st.stop()
    
    return OpenAI(api_key=api_key)

def extract_text_from_docx(file) -> List[Tuple[str, str]]:
    """
    Extract text from Word document, preserving paragraph structure.
    Returns list of tuples: (paragraph_text, paragraph_style)
    """
    doc = Document(file)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():  # Only include non-empty paragraphs
            paragraphs.append((para.text, para.style.name))
    return paragraphs

def extract_text_from_txt(file) -> List[Tuple[str, str]]:
    """
    Extract text from plain text file, preserving paragraph structure.
    Returns list of tuples: (paragraph_text, paragraph_style)
    """
    # Read file content
    content = file.read()
    
    # Try to decode with UTF-8, fallback to other encodings
    if isinstance(content, bytes):
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = content.decode('latin-1')
            except UnicodeDecodeError:
                text = content.decode('utf-8', errors='ignore')
    else:
        text = content
    
    # Split into paragraphs (double newline or single newline for empty lines)
    paragraphs = []
    para_texts = re.split(r'\n\s*\n', text)
    
    for para_text in para_texts:
        para_text = para_text.strip()
        if para_text:  # Only include non-empty paragraphs
            paragraphs.append((para_text, 'Normal'))
    
    # If no double newlines found, split by single newlines
    if not paragraphs:
        lines = text.split('\n')
        current_para = []
        for line in lines:
            line = line.strip()
            if line:
                current_para.append(line)
            else:
                if current_para:
                    paragraphs.append((' '.join(current_para), 'Normal'))
                    current_para = []
        if current_para:
            paragraphs.append((' '.join(current_para), 'Normal'))
    
    return paragraphs

def detect_scanned_pdf(extracted_text: str) -> bool:
    """
    Detect if PDF is likely scanned (image-based).
    Returns True if likely scanned.
    """
    # If extracted text is very short or empty, likely scanned
    text_length = len(extracted_text.strip())
    if text_length < 100:
        return True
    return False

def detect_multi_column_issue(text: str, page_count: int) -> bool:
    """
    Detect potential multi-column layout issues in PDF.
    Returns True if multi-column layout is suspected.
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return False
    
    avg_line_length = sum(len(line) for line in lines) / len(lines)
    
    # Heuristic: short average line length + multiple pages = possible multi-column
    if avg_line_length < 50 and page_count > 1:
        return True
    
    # Additional check: if many lines are very short, likely multi-column
    short_lines = sum(1 for line in lines if len(line) < 40)
    if short_lines > len(lines) * 0.6 and page_count > 1:
        return True
    
    return False

def extract_text_from_pdf_simple(page) -> str:
    """
    Simple text extraction from PDF page using pymupdf.
    Extracts text only (images are automatically excluded).
    """
    try:
        return page.get_text()
    except:
        return ""

def extract_text_from_pdf(file) -> Tuple[List[Tuple[str, str]], dict]:
    """
    Extract text from PDF document using pymupdf, preserving paragraph structure.
    Returns tuple: (paragraphs, metadata)
    paragraphs: list of tuples (paragraph_text, paragraph_style)
    metadata: dict with warnings and info (scanned, multi_column, page_count)
    """
    metadata = {
        'scanned': False,
        'multi_column': False,
        'page_count': 0,
        'encrypted': False,
        'error': None
    }
    
    try:
        # Read file content
        file_content = file.read()
        file.seek(0)  # Reset file pointer
        
        # Open PDF
        doc = fitz.open(stream=file_content, filetype="pdf")
        metadata['page_count'] = len(doc)
        
        # Check if encrypted
        if doc.needs_pass:
            metadata['encrypted'] = True
            doc.close()
            raise ValueError("PDF is password-protected. Please unlock the PDF first.")
        
        # Extract text page by page
        full_text = ""
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = extract_text_from_pdf_simple(page)
            if page_text:
                full_text += page_text + "\n\n"
        
        doc.close()
        
        # Check for scanned PDF
        if detect_scanned_pdf(full_text):
            metadata['scanned'] = True
        
        # Check for multi-column issue (heuristic check)
        if detect_multi_column_issue(full_text, metadata['page_count']):
            metadata['multi_column'] = True
        
        # Split into paragraphs
        paragraphs = []
        if full_text.strip():
            para_texts = re.split(r'\n\s*\n', full_text)
            for para_text in para_texts:
                para_text = para_text.strip()
                if para_text:
                    paragraphs.append((para_text, 'Normal'))
        
        return paragraphs, metadata
        
    except Exception as e:
        metadata['error'] = str(e)
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypted" in error_msg:
            raise ValueError("PDF is password-protected. Please unlock the PDF first.")
        elif "invalid" in error_msg or "corrupted" in error_msg:
            raise ValueError("PDF appears to be corrupted or invalid. Please check the file.")
        else:
            raise ValueError(f"Error reading PDF: {str(e)}")

def extract_text_from_file(file, file_type: str) -> Tuple[List[Tuple[str, str]], Optional[dict]]:
    """
    Unified function to extract text from different file types.
    Returns tuple: (paragraphs, metadata)
    metadata is None for non-PDF files, dict for PDF files
    """
    file.seek(0)  # Reset file pointer
    
    if file_type == 'docx':
        return extract_text_from_docx(file), None
    elif file_type == 'txt':
        return extract_text_from_txt(file), None
    elif file_type == 'pdf':
        return extract_text_from_pdf(file)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def split_paragraphs_into_chunks(paragraphs: List[Tuple[str, str]], max_words: int = 3500) -> List[Tuple[List[Tuple[str, str]], int]]:
    """
    Split paragraphs into chunks that respect sentence and paragraph boundaries.
    Never splits sentences or paragraphs.
    Returns list of tuples: (chunk_paragraphs, chunk_start_index)
    """
    chunks = []
    current_chunk = []
    current_word_count = 0
    chunk_start_idx = 0
    
    for idx, (para_text, para_style) in enumerate(paragraphs):
        para_words = len(para_text.split())
        
        # If a single paragraph exceeds max_words, split by sentences
        if para_words > max_words:
            # Save current chunk if it has content
            if current_chunk:
                chunks.append((current_chunk, chunk_start_idx))
                current_chunk = []
                current_word_count = 0
                chunk_start_idx = idx
            
            # Split paragraph by sentences
            sentences = re.split(r'(?<=[.!?])\s+', para_text)
            sentence_chunk = []
            sentence_word_count = 0
            sentence_start_idx = idx
            
            for sentence in sentences:
                sentence_words = len(sentence.split())
                
                if sentence_word_count + sentence_words > max_words and sentence_chunk:
                    # Save sentence chunk as a special paragraph
                    chunks.append(([(' '.join(sentence_chunk), para_style)], sentence_start_idx))
                    sentence_chunk = [sentence]
                    sentence_word_count = sentence_words
                    sentence_start_idx = idx
                else:
                    sentence_chunk.append(sentence)
                    sentence_word_count += sentence_words
            
            if sentence_chunk:
                current_chunk.append((' '.join(sentence_chunk), para_style))
                current_word_count = sentence_word_count
                if not current_chunk or len(current_chunk) == 1:
                    chunk_start_idx = sentence_start_idx
        else:
            # Check if adding this paragraph would exceed max_words
            if current_word_count + para_words > max_words and current_chunk:
                chunks.append((current_chunk, chunk_start_idx))
                current_chunk = [(para_text, para_style)]
                current_word_count = para_words
                chunk_start_idx = idx
            else:
                current_chunk.append((para_text, para_style))
                current_word_count += para_words
                if len(current_chunk) == 1:
                    chunk_start_idx = idx
    
    # Add remaining chunk
    if current_chunk:
        chunks.append((current_chunk, chunk_start_idx))
    
    return chunks

def translate_text(client: OpenAI, text: str, target_language: str = None, source_language: str = None, translation_hints: str = None) -> str:
    """
    Translate text using OpenAI API.
    """
    system_prompt = "You are a professional translator. Translate the following text accurately while preserving the original formatting, style, and meaning. Maintain paragraph breaks and sentence structure."
    
    user_prompt = text
    if source_language and target_language:
        user_prompt = f"Translate the following text from {source_language} to {target_language}:\n\n{text}"
    elif target_language:
        user_prompt = f"Translate the following text to {target_language}:\n\n{text}"
    elif source_language:
        user_prompt = f"The following text is in {source_language}. Translate it:\n\n{text}"
    
    # Add translation hints if provided
    if translation_hints and translation_hints.strip():
        system_prompt += f"\n\nAdditional translation instructions: {translation_hints}"
    
    try:
        response = client.chat.completions.create(
            model=Config.MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Translation error: {str(e)}")
        return None

def create_docx_from_paragraphs(paragraphs: List[Tuple[str, str]], output_path: str):
    """
    Create a Word document from paragraphs with their original styles.
    """
    doc = Document()
    for text, style_name in paragraphs:
        para = doc.add_paragraph(text)
        try:
            para.style = style_name
        except:
            # If style doesn't exist, use default
            para.style = doc.styles['Normal']
    doc.save(output_path)

def main():
    st.title("🌍 Global Translator: AI Translation for Large Text Files")
    
    # Beautiful explanation section
    st.markdown("""
    <div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 2rem; border-radius: 10px; margin: 1.5rem 0;'>
        <h2 style='color: white; margin-top: 0;'>✨ Professional Document Translation Made Simple</h2>
        <p style='color: white; font-size: 1.1rem; line-height: 1.6; margin-bottom: 0;'>
            Translate large documents with <strong>AI-powered precision</strong> while preserving paragraph structure 
            and sentence boundaries.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Features in columns
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div style='text-align: center; padding: 1rem;'>
            <h3 style='color: #667eea;'>📄 Multi-Format Support</h3>
            <p>Word, Text, and PDF files <small>(text-based PDFs only)</small></p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style='text-align: center; padding: 1rem;'>
            <h3 style='color: #667eea;'>🧠 Smart Chunking & Stitching</h3>
            <p>Intelligent splitting and seamless reassembly</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div style='text-align: center; padding: 1rem;'>
            <h3 style='color: #667eea;'>🌐 50+ Languages</h3>
            <p>Professional translations in any language</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("Settings")
        
        # API Key input (only in prod mode)
        if IS_PROD_MODE:
            st.markdown("### 🔑 API Configuration")
            api_key_input = st.text_input(
                "OpenAI API Key",
                value=st.session_state.get('openai_api_key', ''),
                type="password",
                help="Enter your OpenAI API key. Get one at https://platform.openai.com/api-keys",
                key="api_key_input"
            )
            # Update session state when API key changes
            if api_key_input != st.session_state.get('openai_api_key', ''):
                st.session_state.openai_api_key = api_key_input
            
            if not api_key_input or api_key_input.strip() == '':
                st.warning("⚠️ API key required to use the translator")
            else:
                st.success("✅ API key configured")
            
            st.markdown("---")
        
        source_language = st.selectbox(
            "Source Language (Optional)",
            options=["Auto-detect"] + SUPPORTED_LANGUAGES,
            index=0,
            help="Language of the source document. Select 'Auto-detect' to let the model identify the language automatically."
        )
        # Convert "Auto-detect" to None for the translation function
        source_language_value = None if source_language == "Auto-detect" else source_language
        
        # Find default target language index
        default_target_index = 0
        if Config.DEFAULT_TARGET_LANGUAGE in SUPPORTED_LANGUAGES:
            default_target_index = SUPPORTED_LANGUAGES.index(Config.DEFAULT_TARGET_LANGUAGE)
        
        target_language = st.selectbox(
            "Target Language",
            options=SUPPORTED_LANGUAGES,
            index=default_target_index,
            help="Language to translate the document to"
        )
        chunk_size = st.slider(
            "Chunk Size (words)",
            min_value=2000,
            max_value=5000,
            value=Config.DEFAULT_CHUNK_SIZE,
            step=100,
            help="Number of words per translation chunk"
        )
        
        st.markdown("---")
        st.header("Instructions (Optional)")
        
        # Expandable help section
        with st.expander("ℹ️ How to use Instructions", expanded=False):
            st.markdown("""
            **Provide additional context to improve translation quality:**
            
            **Document Type/Context:**
            - "This is a legal contract - use formal legal terminology"
            - "This is a medical document - preserve medical terms accurately"
            - "This is a technical manual - keep technical terms in English"
            - "This is a marketing brochure - use persuasive, engaging language"
            
            **Style & Tone:**
            - "Use formal language throughout"
            - "Use informal, conversational tone"
            - "Maintain an academic writing style"
            - "Use business/professional language"
            
            **Script/Format Preferences:**
            - "Use Kazakh Latin script" (instead of Cyrillic)
            - "Use Kazakh Cyrillic script"
            - "Use Simplified Chinese" (vs Traditional)
            - "Use British English spelling" (vs American)
            
            **Terminology Handling:**
            - "Preserve all company names, product names, and acronyms in their original form"
            - "Translate technical terms to their standard equivalents"
            - "Keep all proper nouns (names, places) in their original form"
            
            **Special Requirements:**
            - "Preserve all numbers, dates, and percentages exactly as written"
            - "Maintain the original paragraph structure and formatting"
            - "This document contains code snippets - do not translate code, only comments"
            """)
        
        translation_hints = st.text_area(
            "Translation Instructions",
            value="",
            height=100,
            help="Provide additional context or instructions to improve translation quality. Click 'How to use Instructions' above for examples."
        )
    
    # File upload
    uploaded_file = st.file_uploader(
        "Upload Document",
        type=['docx', 'txt', 'pdf'],
        help="Upload a Word (.docx), Text (.txt), or PDF (.pdf) file to translate. Text-based PDFs work best."
    )
    
    # Info about supported formats
    if uploaded_file is None:
        st.info("📄 **Supported formats:** Word (.docx), Text (.txt), PDF (.pdf)\n\n"
                "💡 **Tips:**\n"
                "- Text-based PDFs work best (not scanned)\n"
                "- Single-column PDFs are recommended\n"
                "- Scanned PDFs require OCR (not supported)")
    
    if uploaded_file is not None:
        # Detect file type
        file_name = uploaded_file.name.lower()
        if file_name.endswith('.docx'):
            file_type = 'docx'
            file_type_display = 'Word'
        elif file_name.endswith('.txt'):
            file_type = 'txt'
            file_type_display = 'Text'
        elif file_name.endswith('.pdf'):
            file_type = 'pdf'
            file_type_display = 'PDF'
        else:
            file_type = 'docx'  # Default
            file_type_display = 'Document'
        
        # Display file info
        st.success(f"File uploaded: {uploaded_file.name} ({file_type_display})")
        
        # Process button
        if st.button("Translate Document", type="primary"):
            with st.spinner("Processing document..."):
                try:
                    # Extract text from document
                    paragraphs, metadata = extract_text_from_file(uploaded_file, file_type)
                    
                    # Show PDF-specific warnings
                    if metadata:
                        if metadata.get('encrypted'):
                            st.error("❌ This PDF is password-protected. Please unlock it first.")
                            return
                        
                        if metadata.get('scanned'):
                            st.warning("⚠️ **Scanned PDF Detected:** This PDF appears to be image-based (scanned). "
                                     "Text extraction may be limited. Text-based PDFs work best. "
                                     "Consider using OCR software first.")
                        
                        if metadata.get('multi_column'):
                            st.warning("⚠️ **Multi-Column Layout Detected:** This PDF appears to have multiple columns. "
                                     "The text may be extracted in an incorrect reading order (e.g., all left column text, "
                                     "then all right column text, instead of reading left-to-right across columns). "
                                     "For best results, use single-column PDFs or manually verify the translation output.")
                    
                    if not paragraphs:
                        st.warning("The document appears to be empty or no text could be extracted.")
                        if metadata and metadata.get('scanned'):
                            st.info("💡 Tip: If this is a scanned PDF, you may need OCR software to extract text first.")
                        return
                    
                    # Split paragraphs into chunks (preserving structure)
                    chunk_data = split_paragraphs_into_chunks(paragraphs, max_words=chunk_size)
                    
                    st.info(f"Document split into {len(chunk_data)} chunk(s) for translation.")
                    
                    # Translate chunks
                    client = get_openai_client()
                    translated_paragraphs = []
                    
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    for i, (chunk_paragraphs, start_idx) in enumerate(chunk_data):
                        status_text.text(f"Translating chunk {i+1}/{len(chunk_data)}...")
                        
                        # Combine chunk paragraphs into text for translation
                        chunk_text = '\n\n'.join([para[0] for para in chunk_paragraphs])
                        
                        # Translate the chunk
                        translated_text = translate_text(client, chunk_text, target_language, source_language_value, translation_hints)
                        
                        if translated_text is None:
                            st.error("Translation failed. Please check your API key and try again.")
                            return
                        
                        # Split translated text back into paragraphs
                        # The translation should preserve paragraph breaks
                        translated_para_texts = translated_text.split('\n\n')
                        
                        # Map translated paragraphs back to original styles
                        for j, translated_para_text in enumerate(translated_para_texts):
                            # Use the style from the corresponding original paragraph
                            if j < len(chunk_paragraphs):
                                original_style = chunk_paragraphs[j][1]
                            else:
                                # If we got more paragraphs than expected, use the last style
                                original_style = chunk_paragraphs[-1][1] if chunk_paragraphs else 'Normal'
                            
                            translated_paragraphs.append((translated_para_text.strip(), original_style))
                        
                        progress_bar.progress((i + 1) / len(chunk_data))
                    
                    # Create output document
                    output_paragraphs = translated_paragraphs
                    
                    # Determine output filename and mime type
                    # Always output as DOCX for consistency
                    original_name = uploaded_file.name
                    base_name = os.path.splitext(original_name)[0]
                    output_filename = f"translated_{base_name}.docx"
                    
                    create_docx_from_paragraphs(output_paragraphs, output_filename)
                    
                    # Provide download button
                    with open(output_filename, "rb") as file:
                        st.success("Translation complete!")
                        st.download_button(
                            label="Download Translated Document",
                            data=file.read(),
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Clean up
                    if os.path.exists(output_filename):
                        os.remove(output_filename)
                    
                    status_text.empty()
                    progress_bar.empty()
                    
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
                    st.exception(e)

if __name__ == "__main__":
    main()

